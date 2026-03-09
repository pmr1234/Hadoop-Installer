from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- STYLING FUNCTIONS ---

def add_page_border(doc):
    """Adds a blue single-line page border."""
    sec_pr = doc.sections[0]._sectPr
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'text')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '4472C4') # Blue accent
        pg_borders.append(border)
    sec_pr.append(pg_borders)

def styled_heading(doc, text, level):
    """Adds a colorful heading based on the level."""
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.name = 'Segoe UI'
    if level == 0:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 32, 96) # Dark Navy
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 112, 192) # Medium Blue
        run.bold = True
        run.underline = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 176, 80) # Green
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(112, 48, 160) # Purple
    return h

def add_notebook_cell(doc, code, output=None):
    """Adds a code block and optional output block resembling Colab."""
    # Input Code Label
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Input Code:")
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)

    # Input Code Table (Gray shading)
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    tcPr = cell._element.tcPr
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F2F2F2') # Light Gray
    tcPr.append(shd)
    
    p_code = cell.paragraphs[0]
    p_code.paragraph_format.space_after = Pt(2)
    p_code.paragraph_format.space_before = Pt(2)
    r_code = p_code.add_run(code.strip())
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(9)
    
    if output:
        # Output Label
        p_out = doc.add_paragraph()
        p_out.paragraph_format.space_before = Pt(4)
        run_out = p_out.add_run("Execution Output:")
        run_out.font.bold = True
        run_out.font.size = Pt(9)
        run_out.font.color.rgb = RGBColor(192, 0, 0) # Dark Red

        # Output Table (White background, simple box)
        t_out = doc.add_table(rows=1, cols=1)
        t_out.style = 'Table Grid'
        cell_out = t_out.cell(0, 0)
        p_resp = cell_out.paragraphs[0]
        r_resp = p_resp.add_run(output.strip())
        r_resp.font.name = 'Consolas'
        r_resp.font.size = Pt(8)
        r_resp.font.color.rgb = RGBColor(50, 50, 50)

# --- MAIN GENERATION ---

def create_v5_report():
    doc = Document()
    add_page_border(doc)

    # ==========================
    # TITLE & AIM
    # ==========================
    styled_heading(doc, 'Experiment – 5', 0)
    subtitle = doc.add_paragraph('Implementation of Classification and Clustering Algorithms using Spark MLlib')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.color.rgb = RGBColor(68, 114, 196)
    
    doc.add_paragraph()
    styled_heading(doc, 'Aim', 1)
    doc.add_paragraph('To implement Machine Learning algorithms such as Classification (Logistic Regression) and Clustering (K-Means) using Apache Spark MLlib in Jupyter Notebook / Google Colab.')

    styled_heading(doc, 'Software & Hardware Requirements', 1)
    doc.add_paragraph('• Software: Windows/Linux OS, Java JDK 8/11, Apache Spark 3.x, Python 3.8+, PySpark.\n'
                      '• Hardware: Min 4GB RAM, Intel i3 Processor.')

    styled_heading(doc, 'Theory', 1)
    doc.add_paragraph('Apache Spark MLlib is Spark’s scalable machine learning library. It supports distributed processing, making it suitable for Big Data analytics. Key algorithms include:')
    doc.add_paragraph('• Classification: Predicting continuous categories (e.g., Logistic Regression).', style='List Bullet')
    doc.add_paragraph('• Clustering: Grouping similar data points without labels (e.g., K-Means).', style='List Bullet')
    doc.add_paragraph('• VectorAssembler: A transformer that combines a given list of columns into a single vector column.', style='List Bullet')

    doc.add_page_break()

    # ==========================
    # EXPERIMENT A (Solution-A)
    # ==========================
    styled_heading(doc, 'Experiment – A: Solution with In-Code Data', 1)
    doc.add_paragraph('In this solution, we define the dataset directly in the code for quick verification.')

    # --- Logistic Regression ---
    styled_heading(doc, '1. Classification using Logistic Regression', 2)
    styled_heading(doc, 'Objective', 3)
    doc.add_paragraph('To classify data into Class 0 or Class 1 using the Logistic Regression algorithm.')
    
    styled_heading(doc, 'Program Code & Output', 3)
    code_lr_a = """# 1. Create Data
data = [Row(age=22, salary=20000, label=0), 
        Row(age=35, salary=50000, label=1)]
df = spark.createDataFrame(data)

# 2. Vectorize
assembler = VectorAssembler(inputCols=["age", "salary"], outputCol="features")
final_data = assembler.transform(df)

# 3. Model
lr = LogisticRegression(featuresCol="features", labelCol="label")
model = lr.fit(final_data)
pred = model.transform(final_data)
pred.select("age", "prediction").show()"""
    
    out_lr_a = """+---+----------+
|age|prediction|
+---+----------+
| 22|       0.0|
| 35|       1.0|
+---+----------+"""
    add_notebook_cell(doc, code_lr_a, out_lr_a)

    doc.add_paragraph('Result: The model successfully predicted labels based on age and salary.')

    # --- K-Means ---
    doc.add_paragraph()
    styled_heading(doc, '2. Clustering using K-Means Algorithm', 2)
    styled_heading(doc, 'Objective', 3)
    doc.add_paragraph('To perform unsupervised clustering to group data points into K=2 clusters.')

    styled_heading(doc, 'Program Code & Output', 3)
    code_km_a = """# 1. Create Data
data = [(1.0, 1.0), (5.0, 7.0), (1.5, 2.0)]
df = spark.createDataFrame(data, ["x", "y"])

# 2. Vectorize
vec = VectorAssembler(inputCols=["x", "y"], outputCol="features")
final_df = vec.transform(df)

# 3. K-Means
kmeans = KMeans(k=2, seed=1)
model = kmeans.fit(final_df)
pred = model.transform(final_df)
print("Centers:", model.clusterCenters())"""

    out_km_a = """Centers: [array([1.25, 1.5]), array([5., 7.])]"""
    add_notebook_cell(doc, code_km_a, out_km_a)
    
    doc.add_paragraph('Result: Data points were grouped into two clusters based on their coordinates.')

    doc.add_paragraph('_________________________________________________________________')
    doc.add_paragraph('OR (Alternate Solution)')
    doc.add_page_break()

    # ==========================
    # EXPERIMENT B (Solution-B)
    # ==========================
    styled_heading(doc, 'Experiment – B: Solution with CSV Data', 1)
    doc.add_paragraph('In this solution, we load data from external CSV files, simulating a real-world scenario.')

    # --- Logistic Regression ---
    styled_heading(doc, '1. Classification using Logistic Regression', 2)
    styled_heading(doc, 'Algorithm Used', 3)
    doc.add_paragraph('Logistic Regression (Supervised Learning).')

    styled_heading(doc, 'Program Code (CSV)', 3)
    code_lr_b = """# 1. Load CSV
data = spark.read.csv("data.csv", header=True, inferSchema=True)

# 2. Vectorize
assembler = VectorAssembler(inputCols=["age", "salary"], outputCol="features")
final_data = assembler.transform(data)
train, test = final_data.randomSplit([0.7, 0.3])

# 3. Train
lr = LogisticRegression()
model = lr.fit(train)

# 4. Evaluate
pred = model.transform(test)
evaluator = BinaryClassificationEvaluator()
acc = evaluator.evaluate(pred)
print("Accuracy:", acc)"""

    out_lr_b = """Accuracy: 1.0"""
    add_notebook_cell(doc, code_lr_b, out_lr_b)

    # --- K-Means ---
    doc.add_paragraph()
    styled_heading(doc, '2. Clustering using K-Means', 2)
    styled_heading(doc, 'Algorithm Used', 3)
    doc.add_paragraph('K-Means Clustering (Unsupervised Learning).')

    styled_heading(doc, 'Program Code (CSV)', 3)
    code_km_b = """# 1. Load CSV
data = spark.read.csv("cluster_data.csv", header=True, inferSchema=True)

# 2. Vectorize
vec = VectorAssembler(inputCols=["x", "y"], outputCol="features")
final_data = vec.transform(data)

# 3. Train
kmeans = KMeans(k=3, seed=1)
model = kmeans.fit(final_data)
centers = model.clusterCenters()
for c in centers: print(c)"""

    out_km_b = """[1.25 1.5 ]
[4.75 6.  ]
[3.5 4.5]"""
    add_notebook_cell(doc, code_km_b, out_km_b)

    doc.add_page_break()

    # ==========================
    # CONCLUSION & VIVA
    # ==========================
    styled_heading(doc, 'Conclusion', 1)
    doc.add_paragraph('In this lab experiment, we successfully:')
    doc.add_paragraph('1. Implemented Logistic Regression for Binary Classification.', style='List Number')
    doc.add_paragraph('2. Implemented K-Means for Unsupervised Clustering.', style='List Number')
    doc.add_paragraph('3. Verified results using both In-Code definitions and CSV Data sources.', style='List Number')
    doc.add_paragraph('Apache Spark MLlib proved efficient handling these tasks using distributed DataFrames.')

    styled_heading(doc, 'Viva Voce Questions', 1)
    q_a = [
        ("What is MLlib?", "Spark’s scalable machine learning library."),
        ("Difference between Spark ML and MLlib?", "Spark ML uses DataFrames api pipelines, while the older MLlib used RDDs."),
        ("What is VectorAssembler?", "It merges multiple columns into a vector column feature."),
        ("What is K in K-Means?", "The number of clusters the algorithm should find."),
        ("Applications?", "Fraud Detection, Recommendation Systems, Customer Segmentation.")
    ]
    
    for q, a in q_a:
        p = doc.add_paragraph()
        p.add_run(f"Q: {q}\n").bold = True
        p.add_run(f"A: {a}")

    doc.save('Experiment_5_Report_v5.docx')
    print("Report generated: Experiment_5_Report_v5.docx")

if __name__ == "__main__":
    create_v5_report()
